"""
SEMAI Analytics Intelligence Platform - Export Utilities.

Word-document generation, markdown table parsing, and ZIP file processing.
No Streamlit dependency.
"""

from __future__ import annotations

import zipfile
from datetime import date
from io import BytesIO

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# =============================================================================
# Markdown → HTML table parser (used by the Streamlit renderer)
# =============================================================================

def parse_markdown_table(lines: list[str], start_idx: int):
    """Parse a markdown table starting at *start_idx*.

    Args:
        lines: All lines of the report.
        start_idx: Index of the first table line.

    Returns:
        Tuple of ``(html_string, next_line_index)`` or ``(None, start_idx)``
        if parsing fails.
    """
    table_lines: list[str] = []
    idx = start_idx

    while idx < len(lines) and "|" in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1

    if len(table_lines) < 2:
        return None, start_idx

    headers = [h.strip() for h in table_lines[0].split("|") if h.strip()]

    rows: list[list[str]] = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            rows.append(cells)

    html = '<table class="markdown-table">\n'
    html += "<thead><tr>"
    for header in headers:
        html += f"<th>{header}</th>"
    html += "</tr></thead>\n<tbody>\n"

    for row in rows:
        html += "<tr>"
        for cell in row:
            html += f"<td>{cell}</td>"
        html += "</tr>\n"

    html += "</tbody></table>"
    return html, idx


# =============================================================================
# Word Document Builder
# =============================================================================

def create_word_document(
    report: str,
    site_url: str,
    start_date,
    end_date,
    report_type: str = "Deep Audit",
    period2_start=None,
    period2_end=None,
) -> BytesIO | None:
    """Convert a markdown *report* into a formatted Word document.

    Args:
        report: Markdown text of the report.
        site_url: Property URL or label.
        start_date: Start date (``date`` object).
        end_date: End date (``date`` object).
        report_type: Title prefix for the document.
        period2_start: Optional second-period start date.
        period2_end: Optional second-period end date.

    Returns:
        A ``BytesIO`` buffer containing the ``.docx`` file, or ``None``
        if ``python-docx`` is not installed.
    """
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading(f"{report_type} Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(79, 70, 229)

    # Metadata
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Property: ").bold = True
    meta.add_run(f"{site_url}\n")

    if period2_start and period2_end:
        meta.add_run("Period 1: ").bold = True
        meta.add_run(
            f'{start_date.strftime("%b %d, %Y")} - '
            f'{end_date.strftime("%b %d, %Y")}\n'
        )
        meta.add_run("Period 2: ").bold = True
        meta.add_run(
            f'{period2_start.strftime("%b %d, %Y")} - '
            f'{period2_end.strftime("%b %d, %Y")}\n'
        )
    else:
        meta.add_run("Analysis Period: ").bold = True
        meta.add_run(
            f'{start_date.strftime("%b %d, %Y")} - '
            f'{end_date.strftime("%b %d, %Y")}\n'
        )

    meta.add_run("Generated On: ").bold = True
    meta.add_run(date.today().strftime("%B %d, %Y"))

    doc.add_paragraph()
    doc.add_paragraph("_" * 50)
    doc.add_paragraph()

    # Parse markdown body
    lines = report.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # --- Headers ---
        if line.startswith("# "):
            heading = doc.add_heading(line[2:], 1)
            heading.runs[0].font.color.rgb = RGBColor(79, 70, 229)
        elif line.startswith("## "):
            heading = doc.add_heading(line[3:], 2)
            heading.runs[0].font.color.rgb = RGBColor(79, 70, 229)
        elif line.startswith("### "):
            heading = doc.add_heading(line[4:], 3)
            heading.runs[0].font.color.rgb = RGBColor(67, 56, 202)
        elif line.startswith("#### "):
            heading = doc.add_heading(line[5:], 4)
            heading.runs[0].font.color.rgb = RGBColor(99, 102, 241)

        # --- Tables ---
        elif "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
            table_lines: list[str] = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                headers = [
                    h.strip() for h in table_lines[0].split("|") if h.strip()
                ]
                rows_data: list[list[str]] = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split("|") if c.strip()]
                    if cells:
                        rows_data.append(cells)

                if headers and rows_data:
                    table = doc.add_table(
                        rows=1 + len(rows_data), cols=len(headers)
                    )
                    table.style = "Light Grid Accent 1"

                    for idx, header in enumerate(headers):
                        cell = table.rows[0].cells[idx]
                        cell.text = header
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(
                            79, 70, 229
                        )

                    for row_idx, row_data in enumerate(rows_data, 1):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(headers):
                                table.rows[row_idx].cells[col_idx].text = (
                                    cell_data
                                )

                    doc.add_paragraph()
            continue

        # --- Lists ---
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.5)

        # --- Bold inline markdown ---
        elif "**" in line:
            p = doc.add_paragraph()
            parts = line.split("**")
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                    run.font.color.rgb = RGBColor(79, 70, 229)
                else:
                    p.add_run(part)

        # --- Regular text ---
        else:
            skip_chars = {"---", "═" * 50, "┌", "┐", "└", "┘"}
            if line not in skip_chars:
                doc.add_paragraph(line)

        i += 1

    # Write to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# ZIP / File Processing
# =============================================================================

def process_uploaded_files(
    zip_file,
    file_list: list[str],
) -> pd.DataFrame | None:
    """Read CSV/Excel files from a ZIP archive and combine them.

    Args:
        zip_file: A file-like object (e.g. Streamlit ``UploadedFile``).
        file_list: List of file paths inside the ZIP to process.

    Returns:
        A combined ``DataFrame``, or ``None`` on failure.
    """
    all_dataframes: list[pd.DataFrame] = []

    try:
        with zipfile.ZipFile(zip_file, "r") as zip_ref:
            for filename in file_list:
                try:
                    with zip_ref.open(filename) as fh:
                        file_content = BytesIO(fh.read())
                        if filename.endswith(".csv"):
                            df = pd.read_csv(file_content)
                        else:
                            df = pd.read_excel(file_content)
                        df["source_file"] = filename
                        all_dataframes.append(df)
                except Exception:
                    continue
    except Exception:
        return None

    if not all_dataframes:
        return None

    return pd.concat(all_dataframes, ignore_index=True)


def process_direct_files(uploaded_files) -> "pd.DataFrame | None":
    """Read CSV/Excel files directly from Streamlit UploadedFile objects and combine them.

    Args:
        uploaded_files: List of Streamlit UploadedFile objects.

    Returns:
        A combined ``DataFrame``, or ``None`` on failure.
    """
    all_dataframes: list[pd.DataFrame] = []

    for f in uploaded_files:
        try:
            file_content = BytesIO(f.read())
            if f.name.lower().endswith(".csv"):
                df = pd.read_csv(file_content)
            else:
                df = pd.read_excel(file_content)
            df["source_file"] = f.name
            all_dataframes.append(df)
        except Exception:
            continue

    if not all_dataframes:
        return None

    return pd.concat(all_dataframes, ignore_index=True)
